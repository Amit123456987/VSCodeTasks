import re
import json
from docx import Document

doc = Document("roughCode/distributed.docx")
fileW = open('writeCode.py','w')


para = doc.paragraphs

for row in range(0,len(doc.paragraphs)):
    if(para[row].text[:7] == "Example"):
        
        fileW.write("\npath='"+para[row].text+"'\ncode='''")
        for row2 in range(row+1,len(doc.paragraphs)):
            # print(para[row2].style.font.name)
            if( para[row2].text == "" ):
                fileW.write(para[row2].text)
                continue
            
            elif( para[row2].style.font.size == None ):
                continue
            
            elif(int(para[row2].style.font.size.pt) == 11):
                fileW.write(para[row2].text)
            
            else:
                fileW.write("''' \ncreate_folders_file(path,code)\n")
                break
